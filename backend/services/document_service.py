"""document_service.py - Document creation and file parsing."""
import io
import logging
from typing import List
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete
from fastapi import HTTPException

from backend.models.document_model import Document
from backend.schemas.document_schema import DocumentCreate

logger = logging.getLogger(__name__)


def _extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from PDF, DOCX, or TXT bytes."""
    try:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext == "pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n\n".join(page.get_text() for page in doc)
        elif ext == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext == "txt":
            return file_bytes.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{ext}")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("File text extraction error (%s): %s", filename, exc)
        raise HTTPException(status_code=422, detail=f"Could not extract text from {filename}: {exc}") from exc


async def create_document(db: AsyncSession, user_id: int, data: DocumentCreate) -> Document:
    try:
        doc = Document(
            user_id=user_id,
            name=data.name,
            content=data.content,
            font_size=data.font_size,
            font_style=data.font_style,
            header=data.header,
            footer=data.footer,
        )
        db.add(doc)
        await db.flush()
        await db.refresh(doc)
        logger.info("Document %s created by user %s", doc.id, user_id)
        return doc
    except Exception as exc:
        logger.error("Document creation error: %s", exc)
        raise HTTPException(status_code=500, detail="Document creation failed") from exc


async def create_document_from_file(
    db: AsyncSession, user_id: int, name: str, file_bytes: bytes, filename: str,
    font_size: int = 12, font_style: str = "Helvetica",
    header: str = None, footer: str = None,
) -> Document:
    content = _extract_text_from_bytes(file_bytes, filename)
    try:
        doc = Document(
            user_id=user_id,
            name=name,
            content=content,
            font_size=font_size,
            font_style=font_style,
            header=header,
            footer=footer,
            file_path=filename,
        )
        db.add(doc)
        await db.flush()
        await db.refresh(doc)
        logger.info("Document %s created from file %s by user %s", doc.id, filename, user_id)
        return doc
    except Exception as exc:
        logger.error("Document-from-file creation error: %s", exc)
        raise HTTPException(status_code=500, detail="Document creation from file failed") from exc


async def get_user_documents(db: AsyncSession, user_id: int):
    try:
        result = await db.execute(
            select(Document).where(Document.user_id == user_id).order_by(Document.created_at.desc())
        )
        return result.scalars().all()
    except Exception as exc:
        logger.error("Get documents error for user %s: %s", user_id, exc)
        raise HTTPException(status_code=500, detail="Could not retrieve documents") from exc


async def get_document(db: AsyncSession, doc_id: int, user_id: int) -> Document:
    try:
        result = await db.execute(
            select(Document).where(Document.id == doc_id, Document.user_id == user_id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        return doc
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Get document %s error: %s", doc_id, exc)
        raise HTTPException(status_code=500, detail="Could not retrieve document") from exc

async def delete_document(db: AsyncSession, doc_id: int, user_id: int) -> bool:
    """Permanently delete a document and its associated records via cascade."""
    try:
        doc = await get_document(db, doc_id, user_id)
        await db.delete(doc)
        await db.commit()
        logger.warning("Document %s DELETED by user %s", doc_id, user_id)
        return True
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Delete document %s error: %s", doc_id, exc)
        raise HTTPException(status_code=500, detail="Document deletion failed") from exc

async def bulk_delete_documents(db: AsyncSession, doc_ids: List[int], user_id: int) -> bool:
    """Efficiently delete multiple documents and their associated records via cascade."""
    try:
        stmt = delete(Document).where(Document.id.in_(doc_ids), Document.user_id == user_id)
        await db.execute(stmt)
        await db.commit()
        logger.warning("Bulk DELETED %s documents for user %s", len(doc_ids), user_id)
        return True
    except Exception as exc:
        logger.error("Bulk delete error for user %s: %s", user_id, exc)
        raise HTTPException(status_code=500, detail="Bulk deletion failed") from exc
